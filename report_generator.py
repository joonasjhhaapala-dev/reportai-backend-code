"""
Report Generator - Creates PDF, Word, and Excel reports
"""

import os
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional
import pandas as pd
import openpyxl

# PDF generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT

# Word generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ReportGenerator:
    def __init__(self):
        self.output_dir = Path("outputs")
        self.output_dir.mkdir(exist_ok=True)
    
    async def generate(
        self,
        data: pd.DataFrame,
        workbook: Optional[openpyxl.Workbook],
        config: Dict[str, Any],
        output_format: str = "pdf"
    ) -> Path:
        """
        Generate report in specified format
        
        Args:
            data: Pandas DataFrame with measurement data
            workbook: Optional openpyxl Workbook for Excel files
            config: Report configuration (title, date, company, etc.)
            output_format: Output format (pdf, word, excel)
        
        Returns:
            Path to generated report file
        """
        
        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = config['title'].replace(' ', '_').replace('/', '_')
        filename = f"{safe_title}_{timestamp}.{self._get_extension(output_format)}"
        output_path = self.output_dir / filename
        
        # Generate based on format
        if output_format == "pdf":
            self._generate_pdf(data, config, output_path)
        elif output_format == "word":
            self._generate_word(data, config, output_path)
        elif output_format == "excel":
            self._generate_excel(data, workbook, config, output_path)
        else:
            raise ValueError(f"Unsupported format: {output_format}")
        
        return output_path
    
    def _get_extension(self, format: str) -> str:
        """Get file extension for format"""
        extensions = {
            "pdf": "pdf",
            "word": "docx",
            "excel": "xlsx"
        }
        return extensions.get(format, "pdf")
    
    def _generate_pdf(self, data: pd.DataFrame, config: Dict[str, Any], output_path: Path):
        """Generate PDF report"""
        
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Container for elements
        elements = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#00ff88'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#00ff88'),
            spaceAfter=12
        )
        
        # Title
        title = Paragraph(config['title'], title_style)
        elements.append(title)
        elements.append(Spacer(1, 12))
        
        # Metadata table
        metadata = [
            ['Date:', config['date']],
            ['Company:', config.get('company', 'N/A')],
            ['Author:', config.get('author', 'N/A')],
            ['Generated:', datetime.now().strftime('%Y-%m-%d %H:%M')]
        ]
        
        meta_table = Table(metadata, colWidths=[2*inch, 4*inch])
        meta_table.setStyle(TableStyle([
            ('FONT', (0, 0), (-1, -1), 'Helvetica', 10),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.grey),
            ('TEXTCOLOR', (1, 0), (1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        elements.append(meta_table)
        elements.append(Spacer(1, 20))
        
        # Executive Summary
        analysis = config.get('analysis', {})
        
        elements.append(Paragraph('Executive Summary', heading_style))
        elements.append(Spacer(1, 12))
        summary_text = analysis.get('executive_summary', 'No summary available')
        elements.append(Paragraph(summary_text, styles['BodyText']))
        elements.append(Spacer(1, 20))
        
        # Key Findings
        elements.append(Paragraph('Key Findings', heading_style))
        elements.append(Spacer(1, 12))
        findings = analysis.get('key_findings', [])
        for i, finding in enumerate(findings, 1):
            elements.append(Paragraph(f"{i}. {finding}", styles['BodyText']))
            elements.append(Spacer(1, 6))
        elements.append(Spacer(1, 20))
        
        # Statistical Analysis
        elements.append(Paragraph('Statistical Analysis', heading_style))
        elements.append(Spacer(1, 12))
        stat_analysis = analysis.get('statistical_analysis', 'No statistical analysis available')
        elements.append(Paragraph(stat_analysis, styles['BodyText']))
        elements.append(Spacer(1, 20))
        
        # Data Summary Table
        elements.append(Paragraph('Data Summary', heading_style))
        elements.append(Spacer(1, 12))
        
        # Show first few rows of data
        sample_data = data.head(10)
        table_data = [sample_data.columns.tolist()] + sample_data.values.tolist()
        
        # Limit columns if too many
        if len(table_data[0]) > 6:
            table_data = [row[:6] for row in table_data]
        
        data_table = Table(table_data)
        data_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#00ff88')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#0a0f1a')),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
        ]))
        elements.append(data_table)
        elements.append(Spacer(1, 20))
        
        # Recommendations
        elements.append(PageBreak())
        elements.append(Paragraph('Recommendations', heading_style))
        elements.append(Spacer(1, 12))
        recommendations = analysis.get('recommendations', [])
        for i, rec in enumerate(recommendations, 1):
            elements.append(Paragraph(f"{i}. {rec}", styles['BodyText']))
            elements.append(Spacer(1, 6))
        elements.append(Spacer(1, 20))
        
        # Conclusion
        elements.append(Paragraph('Conclusion', heading_style))
        elements.append(Spacer(1, 12))
        conclusion = analysis.get('conclusion', 'No conclusion available')
        elements.append(Paragraph(conclusion, styles['BodyText']))
        
        # Footer
        elements.append(Spacer(1, 40))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.grey,
            alignment=TA_CENTER
        )
        footer_text = f"Generated by ReportAI - Automated Quality Reports | {datetime.now().strftime('%Y-%m-%d')}"
        elements.append(Paragraph(footer_text, footer_style))
        
        # Build PDF
        doc.build(elements)
    
    def _generate_word(self, data: pd.DataFrame, config: Dict[str, Any], output_path: Path):
        """Generate Word (DOCX) report"""
        
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = config['title']
        doc.core_properties.author = config.get('author', 'ReportAI')
        doc.core_properties.created = datetime.now()
        
        # Title
        title = doc.add_heading(config['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Date: {config['date']}")
        doc.add_paragraph(f"Company: {config.get('company', 'N/A')}")
        doc.add_paragraph(f"Author: {config.get('author', 'N/A')}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph()
        
        # Executive Summary
        analysis = config.get('analysis', {})
        
        doc.add_heading('Executive Summary', 1)
        doc.add_paragraph(analysis.get('executive_summary', 'No summary available'))
        doc.add_paragraph()
        
        # Key Findings
        doc.add_heading('Key Findings', 1)
        findings = analysis.get('key_findings', [])
        for finding in findings:
            doc.add_paragraph(finding, style='List Bullet')
        doc.add_paragraph()
        
        # Statistical Analysis
        doc.add_heading('Statistical Analysis', 1)
        doc.add_paragraph(analysis.get('statistical_analysis', 'No statistical analysis available'))
        doc.add_paragraph()
        
        # Data Summary
        doc.add_heading('Data Summary', 1)
        
        # Add table with data sample
        sample_data = data.head(10)
        table = doc.add_table(rows=len(sample_data)+1, cols=min(len(sample_data.columns), 6))
        table.style = 'Light Grid Accent 1'
        
        # Header row
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(sample_data.columns[:6]):
            hdr_cells[i].text = str(col)
        
        # Data rows
        for i, row in enumerate(sample_data.values):
            row_cells = table.rows[i+1].cells
            for j, value in enumerate(row[:6]):
                row_cells[j].text = str(value)
        
        doc.add_paragraph()
        
        # Recommendations
        doc.add_page_break()
        doc.add_heading('Recommendations', 1)
        recommendations = analysis.get('recommendations', [])
        for rec in recommendations:
            doc.add_paragraph(rec, style='List Number')
        doc.add_paragraph()
        
        # Conclusion
        doc.add_heading('Conclusion', 1)
        doc.add_paragraph(analysis.get('conclusion', 'No conclusion available'))
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph(f"Generated by ReportAI - {datetime.now().strftime('%Y-%m-%d')}")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save
        doc.save(str(output_path))
    
    def _generate_excel(
        self, 
        data: pd.DataFrame, 
        workbook: Optional[openpyxl.Workbook],
        config: Dict[str, Any], 
        output_path: Path
    ):
        """Generate Excel report with original data + analysis"""
        
        # Create new workbook
        wb = openpyxl.Workbook()
        
        # Summary sheet
        ws_summary = wb.active
        ws_summary.title = "Summary"
        
        # Write metadata
        ws_summary['A1'] = config['title']
        ws_summary['A1'].font = openpyxl.styles.Font(size=18, bold=True)
        
        ws_summary['A3'] = 'Date:'
        ws_summary['B3'] = config['date']
        ws_summary['A4'] = 'Company:'
        ws_summary['B4'] = config.get('company', 'N/A')
        ws_summary['A5'] = 'Author:'
        ws_summary['B5'] = config.get('author', 'N/A')
        
        # Analysis
        analysis = config.get('analysis', {})
        
        ws_summary['A7'] = 'Executive Summary'
        ws_summary['A7'].font = openpyxl.styles.Font(size=14, bold=True)
        ws_summary['A8'] = analysis.get('executive_summary', '')
        ws_summary['A8'].alignment = openpyxl.styles.Alignment(wrap_text=True)
        
        # Key Findings
        ws_summary['A10'] = 'Key Findings'
        ws_summary['A10'].font = openpyxl.styles.Font(size=14, bold=True)
        findings = analysis.get('key_findings', [])
        for i, finding in enumerate(findings, 11):
            ws_summary[f'A{i}'] = f"• {finding}"
        
        # Data sheet
        ws_data = wb.create_sheet("Data")
        
        # Write DataFrame to Excel
        for r_idx, row in enumerate(data.itertuples(index=False), 1):
            for c_idx, value in enumerate(row, 1):
                ws_data.cell(row=r_idx+1, column=c_idx, value=value)
        
        # Write headers
        for c_idx, col_name in enumerate(data.columns, 1):
            cell = ws_data.cell(row=1, column=c_idx, value=col_name)
            cell.font = openpyxl.styles.Font(bold=True)
            cell.fill = openpyxl.styles.PatternFill(start_color="00FF88", end_color="00FF88", fill_type="solid")
        
        # Save workbook
        wb.save(str(output_path))
